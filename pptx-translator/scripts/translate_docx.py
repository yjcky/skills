#!/usr/bin/env python3
"""
DOCX Translator - Translate Word documents using Amazon Translate, Bedrock LLM,
Google Translate, or Cerebras LLM.

Usage:
    python translate_docx.py input.docx --source zh --target en
    python translate_docx.py input.docx --source en --target zh --engine bedrock
    python translate_docx.py input.docx --source zh --target en --engine cerebras
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document

# Import shared translation infrastructure
from translate_common import (
    TranslationEngine, AmazonTranslateEngine, BedrockLLMEngine,
    GoogleTranslateEngine, CerebrasEngine,
    create_engine, load_glossary, import_terminology,
    post_process_translation,
    LANGUAGE_NAMES,
    safe_resolve_path, setup_windows_encoding, add_common_arguments,
    build_engine_from_args,
    token_aware_batch, estimate_tokens, estimate_prompt_overhead,
)


# ---------------------------------------------------------------------------
# DOCX-specific: text extraction
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Formatting-aware text extraction & application
# ---------------------------------------------------------------------------

# Tags used to mark formatting in source text for LLM preservation
# <b> = bold, <i> = italic, <bi> = bold+italic
_FMT_TAG_RE = re.compile(r'(</?bi>|</?b>|</?i>)')


def _build_marked_text(para):
    """Wrap formatted runs in XML-like tags for markup-preserving translation.

    Returns:
        (marked_text, marked_runs_info)
        marked_runs_info: list of {bold, italic, has_text, char_count} per run
    """
    marked_runs_info = []
    parts = []

    for run in para.runs:
        is_bold = bool(run.bold)
        is_italic = bool(run.italic)
        has_text = bool(run.text.strip())
        info = {
            'bold': is_bold,
            'italic': is_italic,
            'has_text': has_text,
            'char_count': len(run.text),
        }
        marked_runs_info.append(info)

        text = run.text
        if not text.strip():
            parts.append(text)
        elif is_bold and is_italic:
            parts.append(f'<bi>{text}</bi>')
        elif is_bold:
            parts.append(f'<b>{text}</b>')
        elif is_italic:
            parts.append(f'<i>{text}</i>')
        else:
            parts.append(text)

    return ''.join(parts), marked_runs_info


def _parse_formatted_translation(translated_text, marked_runs_info):
    """Parse tag-preserving translated text back into per-run strings.

    Tokenizes the translated text at formatting tag boundaries, collects
    consecutive same-format segments, then assigns them to runs in order.
    Same-format runs whose tags were merged by the LLM are split proportionally.

    Returns list of strings (one per run) or None if tags were lost.
    """
    if not _FMT_TAG_RE.search(translated_text):
        return None  # Tags stripped — caller must fall back

    # ── Tokenize at tag boundaries ──
    tokens = _FMT_TAG_RE.split(translated_text)

    bold = italic = False
    segments = []  # [{text, bold, italic}]

    for tok in tokens:
        if not tok:
            continue
        if tok == '<b>':
            bold = True
        elif tok == '</b>':
            bold = False
        elif tok == '<i>':
            italic = True
        elif tok == '</i>':
            italic = False
        elif tok == '<bi>':
            bold = italic = True
        elif tok == '</bi>':
            bold = italic = False
        else:
            seg_text = tok.strip()
            if seg_text:
                segments.append({'text': seg_text, 'bold': bold, 'italic': italic})

    if not segments:
        return None

    # ── Assign segments to runs (format-guided, in order) ──
    result = []
    seg_idx = 0
    seg_count = len(segments)

    for info in marked_runs_info:
        if not info['has_text']:
            result.append('')
            continue

        if seg_idx >= seg_count:
            result.append('')
            continue

        # Collect consecutive segments matching this run's format
        parts = []
        while (seg_idx < seg_count
               and segments[seg_idx]['bold'] == info['bold']
               and segments[seg_idx]['italic'] == info['italic']):
            parts.append(segments[seg_idx]['text'])
            seg_idx += 1

        if parts:
            result.append(''.join(parts))
        else:
            # Format mismatch — take the next segment anyway
            result.append(segments[seg_idx]['text'])
            seg_idx += 1

    # Pad to match run count
    while len(result) < len(marked_runs_info):
        result.append('')

    # ── Handle leftover segments (LLM hallucinated extra tags) ──
    if seg_idx < seg_count:
        # Merge all remaining segment text into the last valid run
        leftover = ''.join(s['text'] for s in segments[seg_idx:])
        for i in range(len(result) - 1, -1, -1):
            if result[i]:
                result[i] += leftover
                break
        else:
            # No non-empty run — put everything in the first text-bearing run
            for i, info in enumerate(marked_runs_info):
                if info['has_text']:
                    if i < len(result):
                        result[i] = leftover
                    break

    # ── Post-process: split merged same-format runs proportionally ──
    final = []
    res_idx = 0
    run_idx = 0

    while run_idx < len(marked_runs_info) and res_idx < len(result):
        info = marked_runs_info[run_idx]
        if not info['has_text']:
            final.append('')
            run_idx += 1
            continue

        # Count consecutive runs with identical formatting
        same_count = 1
        while run_idx + same_count < len(marked_runs_info):
            nxt = marked_runs_info[run_idx + same_count]
            if (nxt['bold'] == info['bold']
                    and nxt['italic'] == info['italic']
                    and nxt['has_text']):
                same_count += 1
            else:
                break

        if same_count == 1:
            final.append(result[res_idx])
            run_idx += 1
        else:
            # Split proportionally among same-format runs
            char_counts = [marked_runs_info[run_idx + j]['char_count']
                           for j in range(same_count)]
            splits = _split_text_proportionally(result[res_idx], char_counts)
            final.extend(splits)
            run_idx += same_count

        res_idx += 1

    while len(final) < len(marked_runs_info):
        final.append('')

    # ── Ensure spacing between adjacent runs of different formats ──
    # Chinese source has no spaces between words, so when translated to English,
    # spaces are needed at format boundaries (e.g. bold→normal, italic→normal)
    for i in range(len(final) - 1):
        if not final[i] or not final[i + 1]:
            continue
        ri = marked_runs_info[i]
        rj = marked_runs_info[i + 1]
        if not ri['has_text'] or not rj['has_text']:
            continue
        if ri['bold'] != rj['bold'] or ri['italic'] != rj['italic']:
            # Different formatting on adjacent runs — ensure a space separates them
            if not final[i].endswith(' ') and not final[i + 1].startswith(' '):
                final[i] += ' '

    return final[:len(marked_runs_info)]


def _split_text_proportionally(translated_text, char_counts):
    """Split translated text across runs proportionally (largest remainder method).

    Used as a fallback when tags are lost, or to split merged same-format runs.
    """
    if not char_counts:
        return [translated_text]

    n_text = len(translated_text)
    if n_text == 0:
        return [''] * len(char_counts)

    total = sum(char_counts)
    if total == 0:
        result = [''] * len(char_counts)
        if result:
            result[0] = translated_text
        return result

    targets = [count / total * n_text for count in char_counts]
    base = [int(t) for t in targets]

    remaining = n_text - sum(base)
    if remaining > 0:
        remainders = [(targets[i] - base[i], i) for i in range(len(base))]
        remainders.sort(key=lambda x: x[0], reverse=True)
        for j in range(remaining):
            base[remainders[j][1]] += 1

    result = []
    pos = 0
    for length in base:
        if length > 0:
            result.append(translated_text[pos:pos + length])
            pos += length
        else:
            result.append('')
    return result


def _merge_paragraph_text(para):
    """Merge all runs into a single string; also return per-run char counts."""
    runs_info = [len(run.text) for run in para.runs]
    merged = ''.join(run.text for run in para.runs if run.text.strip())
    return merged, runs_info


def extract_texts_from_document(doc) -> list:
    """Extract all translatable texts from a Word document with location metadata.

    Covers: body paragraphs, tables, headers, and footers.
    """
    texts = []

    # --- Body paragraphs ---
    for para_idx, para in enumerate(doc.paragraphs):
        text, marked_runs_info = _build_marked_text(para)
        if text.strip():
            texts.append({
                'text': text,
                'location': ('paragraph', para_idx),
                'original_text': text,
                'paragraph': para,
                'marked_runs_info': marked_runs_info,
            })

    # --- Tables ---
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    text, marked_runs_info = _build_marked_text(para)
                    if text.strip():
                        texts.append({
                            'text': text,
                            'location': ('table', table_idx, row_idx, cell_idx, para_idx),
                            'original_text': text,
                            'paragraph': para,
                            'marked_runs_info': marked_runs_info,
                        })

    # --- Headers & Footers ---
    for section_idx, section in enumerate(doc.sections):
        # Header
        if section.header:
            for para_idx, para in enumerate(section.header.paragraphs):
                text, marked_runs_info = _build_marked_text(para)
                if text.strip():
                    texts.append({
                        'text': text,
                        'location': ('header', section_idx, para_idx),
                        'original_text': text,
                        'paragraph': para,
                        'marked_runs_info': marked_runs_info,
                    })
            # Header tables
            for table_idx, table in enumerate(section.header.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            text, marked_runs_info = _build_marked_text(para)
                            if text.strip():
                                texts.append({
                                    'text': text,
                                    'location': ('header_table', section_idx, table_idx, row_idx, cell_idx, para_idx),
                                    'original_text': text,
                                    'paragraph': para,
                                    'marked_runs_info': marked_runs_info,
                                })

        # Footer
        if section.footer:
            for para_idx, para in enumerate(section.footer.paragraphs):
                text, marked_runs_info = _build_marked_text(para)
                if text.strip():
                    texts.append({
                        'text': text,
                        'location': ('footer', section_idx, para_idx),
                        'original_text': text,
                        'paragraph': para,
                        'marked_runs_info': marked_runs_info,
                    })
            # Footer tables
            for table_idx, table in enumerate(section.footer.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            text, marked_runs_info = _build_marked_text(para)
                            if text.strip():
                                texts.append({
                                    'text': text,
                                    'location': ('footer_table', section_idx, table_idx, row_idx, cell_idx, para_idx),
                                    'original_text': text,
                                    'paragraph': para,
                                    'marked_runs_info': marked_runs_info,
                                })

    return texts


# ---------------------------------------------------------------------------
# DOCX-specific: text application
# ---------------------------------------------------------------------------

def find_paragraph_by_location(location, doc):
    """Resolve a location tuple back to the corresponding paragraph object."""
    loc_type = location[0]

    if loc_type == 'paragraph':
        _, para_idx = location
        if para_idx < len(doc.paragraphs):
            return doc.paragraphs[para_idx]

    elif loc_type == 'table':
        _, table_idx, row_idx, cell_idx, para_idx = location
        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                if cell_idx < len(row.cells):
                    cell = row.cells[cell_idx]
                    if para_idx < len(cell.paragraphs):
                        return cell.paragraphs[para_idx]

    elif loc_type == 'header':
        _, section_idx, para_idx = location
        if section_idx < len(doc.sections):
            header = doc.sections[section_idx].header
            if header and para_idx < len(header.paragraphs):
                return header.paragraphs[para_idx]

    elif loc_type == 'header_table':
        _, section_idx, table_idx, row_idx, cell_idx, para_idx = location
        if section_idx < len(doc.sections):
            header = doc.sections[section_idx].header
            if header and table_idx < len(header.tables):
                table = header.tables[table_idx]
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if cell_idx < len(row.cells):
                        cell = row.cells[cell_idx]
                        if para_idx < len(cell.paragraphs):
                            return cell.paragraphs[para_idx]

    elif loc_type == 'footer':
        _, section_idx, para_idx = location
        if section_idx < len(doc.sections):
            footer = doc.sections[section_idx].footer
            if footer and para_idx < len(footer.paragraphs):
                return footer.paragraphs[para_idx]

    elif loc_type == 'footer_table':
        _, section_idx, table_idx, row_idx, cell_idx, para_idx = location
        if section_idx < len(doc.sections):
            footer = doc.sections[section_idx].footer
            if footer and table_idx < len(footer.tables):
                table = footer.tables[table_idx]
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if cell_idx < len(row.cells):
                        cell = row.cells[cell_idx]
                        if para_idx < len(cell.paragraphs):
                            return cell.paragraphs[para_idx]

    return None


def apply_paragraph_translation(para, translated_text, marked_runs_info=None):
    """Apply translated text to a paragraph.

    Strategy (tried in order):
    1. Tag-based parsing: if marked_runs_info is provided, parse formatting
       tags (<b>, <i>, <bi>) from the LLM output to reconstruct per-run texts
       with correct formatting boundaries (no mid-word breaks).
    2. Fallback: if tags were lost, put all text in the first run.
    """
    if not para or not translated_text:
        return

    runs = para.runs
    if not runs:
        para.add_run(translated_text)
        return

    if marked_runs_info:
        # Try tag-based parsing first (best quality — preserves format boundaries)
        result = _parse_formatted_translation(translated_text, marked_runs_info)
        if result is not None:
            for i, run in enumerate(runs):
                if i < len(result):
                    run.text = result[i]
                else:
                    run.text = ''
            return

    # Fallback: first run gets all translated text
    runs[0].text = translated_text
    for run in runs[1:]:
        run.text = ""


# ---------------------------------------------------------------------------
# Sequential translation fallback (for non-batch engines)
# ---------------------------------------------------------------------------

def translate_document_sequential(doc, engine: TranslationEngine, source_lang: str, target_lang: str):
    """Translate the entire document one paragraph at a time (for non-LLM engines)."""
    todo = []

    # Collect body paragraphs
    for para in doc.paragraphs:
        text, _ = _merge_paragraph_text(para)
        if text.strip():
            todo.append(para)

    # Collect table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text, _ = _merge_paragraph_text(para)
                    if text.strip():
                        todo.append(para)

    # Collect headers/footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                text, _ = _merge_paragraph_text(para)
                if text.strip():
                    todo.append(para)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text, _ = _merge_paragraph_text(para)
                            if text.strip():
                                todo.append(para)
        if section.footer:
            for para in section.footer.paragraphs:
                text, _ = _merge_paragraph_text(para)
                if text.strip():
                    todo.append(para)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text, _ = _merge_paragraph_text(para)
                            if text.strip():
                                todo.append(para)

    total = len(todo)
    for idx, para in enumerate(todo, 1):
        print(f"  Paragraph {idx}/{total}")
        translated = False
        for run in para.runs:
            if run.text.strip():
                translated_text = engine.translate(run.text, source_lang, target_lang)
                run.text = post_process_translation(translated_text, target_lang)
                translated = True
        # If no runs were translated but paragraph has text, translate the first run
        if not translated and para.runs:
            para_text, _ = _merge_paragraph_text(para)
            if para_text.strip():
                translated_text = engine.translate(para_text, source_lang, target_lang)
                para.runs[0].text = post_process_translation(translated_text, target_lang)


# ---------------------------------------------------------------------------
# Main translation logic
# ---------------------------------------------------------------------------

def translate_document(doc, engine: TranslationEngine, source_lang: str,
                       target_lang: str, batch_mode: bool = False,
                       auto_batch: bool = False, max_batch_tokens: int = 25000):
    """Translate all text in a Word document."""

    if batch_mode and (isinstance(engine, (BedrockLLMEngine, CerebrasEngine))):
        print("Extracting texts...")
        text_items = extract_texts_from_document(doc)

        if not text_items:
            print("No text found to translate.")
            return

        print(f"Found {len(text_items)} text segments")

        texts = [item['text'] for item in text_items]

        has_glossary = bool(getattr(engine, 'glossary', None))

        if auto_batch:
            batches = token_aware_batch(texts, max_tokens=max_batch_tokens,
                                        target_lang=target_lang, has_glossary=has_glossary)
            batch_count = len(batches)
            # Show batch stats
            batch_sizes = [len(b) for b in batches]
            batch_tokens = []
            batch_fixed = 250 + (50 if has_glossary else 0)
            for b in batches:
                content_tokens = sum(estimate_tokens(t, target_lang) + 8 for t in b)
                batch_tokens.append(batch_fixed + content_tokens)
            avg_tokens = sum(batch_tokens) // max(batch_count, 1)
            print(f"  Auto-batched into {batch_count} batches "
                  f"(sizes: {batch_sizes}, est. tokens avg/max: {avg_tokens}/{max(batch_tokens)})")
        else:
            batch_size = engine.batch_size
            batches = []
            for i in range(0, len(texts), batch_size):
                batches.append(texts[i:i + batch_size])
            batch_count = len(batches)

        import concurrent.futures
        import time
        translated_texts = []

        # Cerebras rate limiting: 5 req/min → need 12-15s between batches
        is_cerebras = isinstance(engine, CerebrasEngine)
        if is_cerebras:
            max_workers = 1
            min_delay = 14.0  # seconds between batches (5/min = 12s, add buffer)
        else:
            max_workers = min(3, batch_count)
            min_delay = 0

        print(f"Translating {batch_count} batches with {max_workers} concurrent workers"
              + (f" (min {min_delay:.0f}s delay for rate limit)" if is_cerebras else "")
              + "...")

        results_by_index = {}
        failed_count = 0
        _last_submit_time = [0.0]  # mutable for closure

        def submit_batch(idx, batch):
            """Submit a batch with rate-limit pacing."""
            if is_cerebras:
                elapsed = time.time() - _last_submit_time[0]
                if elapsed < min_delay:
                    time.sleep(min_delay - elapsed)
            result = engine.translate_batch(batch, source_lang, target_lang)
            if is_cerebras:
                _last_submit_time[0] = time.time()
            return idx, result

        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_idx = {}
            for idx, batch in enumerate(batches):
                future = executor.submit(submit_batch, idx, batch)
                future_to_idx[future] = idx
                if is_cerebras and idx == 0:
                    _last_submit_time[0] = time.time()

            for future in concurrent.futures.as_completed(future_to_idx):
                idx = future_to_idx[future]
                try:
                    bidx, result = future.result()
                    results_by_index[bidx] = result
                except Exception as e:
                    failed_count += 1
                    print(f"  [ERROR] Batch {idx + 1}/{batch_count} failed: {e}")
                    results_by_index[idx] = [str(t) for t in batches[idx]]

            for i in range(batch_count):
                batch_result = results_by_index.get(i, [str(t) for t in batches[i]])
                translated_texts.extend(batch_result)

        if batch_count > 0 and failed_count == batch_count:
            raise RuntimeError(
                f"All {batch_count} translation batches failed. "
                "Check network connectivity to Cerebras API."
            )

        print("Applying translations...")
        if len(text_items) != len(translated_texts):
            print(f"  [WARN] Translation count mismatch: {len(text_items)} original vs {len(translated_texts)} translated")
            while len(translated_texts) < len(text_items):
                original_item = text_items[len(translated_texts)]
                translated_texts.append(str(original_item['text']) if original_item['text'] else "")

        for idx, (item, translated) in enumerate(zip(text_items, translated_texts)):
            if not translated or len(str(translated).strip()) == 0:
                print(f"  [WARN] Empty translation at position {idx}, using original text")
                translated = str(item['text']) if item['text'] else ""

            if isinstance(translated, dict):
                print(f"  [WARN] Dictionary format at position {idx}, converting to string")
                translated_text = str(translated)
            elif isinstance(translated, list):
                print(f"  [WARN] List format at position {idx}, joining to string")
                translated_text = ' '.join(str(t) for t in translated)
            else:
                translated_text = str(translated)

            translated_text = post_process_translation(translated_text, target_lang)

            para = item.get('paragraph')
            if para is not None:
                try:
                    apply_paragraph_translation(para, translated_text, item.get('marked_runs_info'))
                    loc_type = item['location'][0]
                    print(f"  Applied translation to {loc_type} #{idx + 1}")
                except Exception as e:
                    print(f"  [ERROR] Failed to apply translation at position {idx}: {e}")
                    apply_paragraph_translation(para, item['original_text'] if item['original_text'] else "")
            else:
                # Fallback: resolve by location
                para = find_paragraph_by_location(item['location'], doc)
                if para:
                    try:
                        apply_paragraph_translation(para, translated_text)
                        print(f"  Applied translation via location lookup #{idx + 1}")
                    except Exception as e:
                        print(f"  [ERROR] Failed to apply translation at position {idx}: {e}")
                else:
                    print(f"  [ERROR] Could not find paragraph at position {idx} with location {item['location']}")
    else:
        # Sequential mode
        translate_document_sequential(doc, engine, source_lang, target_lang)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Translate DOCX files using Amazon Translate, Bedrock LLM, Google Translate, or Cerebras LLM'
    )
    parser.add_argument('input_file', help='Input DOCX file path')
    add_common_arguments(parser)

    args = parser.parse_args()

    setup_windows_encoding()

    input_path = safe_resolve_path(args.input_file)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input_file}")
        sys.exit(1)

    if args.output:
        output_path = safe_resolve_path(args.output).parent / Path(args.output).name
        output_path.parent.mkdir(parents=True, exist_ok=True)
    else:
        output_path = input_path.with_name(f"{input_path.stem}-{args.target}.docx")

    engine, batch_mode = build_engine_from_args(args)

    print(f"Loading {input_path}...")
    doc = Document(str(input_path))

    print(f"Translating from {args.source} to {args.target}...")
    try:
        translate_document(doc, engine, args.source, args.target, batch_mode=batch_mode,
                            auto_batch=args.auto_batch, max_batch_tokens=args.max_batch_tokens)
    except RuntimeError as e:
        print(f"Error: {e}")
        sys.exit(1)

    print(f"Saving {output_path}...")
    doc.save(str(output_path.resolve()))
    print("Done!")


if __name__ == '__main__':
    main()
